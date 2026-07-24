"""Word（.docx）解析器测试（用 python-docx 动态生成真实 docx）"""

import docx
import pytest

from app.parsers.docx_parser import DOCX_MIME, DocxParser
from app.parsers.registry import ParserRegistry


def _make_docx(path) -> None:
    doc = docx.Document()
    doc.add_heading("公司考勤制度", level=1)
    doc.add_paragraph("第一章 总则")
    doc.add_paragraph("为标准考勤管理，特制定本制度。")
    doc.add_heading("请假流程", level=2)
    doc.add_paragraph("员工请假需提前一天申请。")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "类型"
    table.rows[0].cells[1].text = "天数上限"
    table.rows[0].cells[2].text = "审批人"
    table.rows[1].cells[0].text = "事假"
    table.rows[1].cells[1].text = "3天"
    table.rows[1].cells[2].text = "部门主管"
    doc.save(str(path))


@pytest.mark.unit
class TestDocxParser:
    """真实 docx 解析"""

    def test_parse_paragraphs(self, tmp_path):
        """正文段落完整提取"""
        f = tmp_path / "test.docx"
        _make_docx(f)

        result = DocxParser().parse(str(f))

        assert "第一章 总则" in result
        assert "员工请假需提前一天申请" in result

    def test_parse_headings_with_hash_prefix(self, tmp_path):
        """标题映射为 Markdown # 前缀（与分块器兼容）"""
        f = tmp_path / "test.docx"
        _make_docx(f)

        result = DocxParser().parse(str(f))

        assert "# 公司考勤制度" in result
        assert "## 请假流程" in result

    def test_parse_table_rows(self, tmp_path):
        """表格按行展开，单元格以 | 分隔"""
        f = tmp_path / "test.docx"
        _make_docx(f)

        result = DocxParser().parse(str(f))

        assert "类型 | 天数上限 | 审批人" in result
        assert "事假 | 3天 | 部门主管" in result

    def test_file_not_found(self, tmp_path):
        """文件不存在抛 FileNotFoundError"""
        with pytest.raises(FileNotFoundError, match="文件不存在"):
            DocxParser().parse(str(tmp_path / "missing.docx"))

    def test_invalid_docx_raises_value_error(self, tmp_path):
        """伪 docx（垃圾字节）抛 ValueError"""
        f = tmp_path / "fake.docx"
        f.write_bytes(b"this is not a docx file")

        with pytest.raises(ValueError, match="无法打开"):
            DocxParser().parse(str(f))

    def test_supported_formats(self):
        """声明支持 docx MIME"""
        assert DOCX_MIME in DocxParser().supported_formats

    def test_registry_integration(self):
        """ParserRegistry 能为 docx MIME 返回 DocxParser"""
        parser = ParserRegistry.get_parser(DOCX_MIME)
        assert isinstance(parser, DocxParser)
