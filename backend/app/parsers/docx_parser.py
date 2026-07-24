"""
Word（.docx）解析器

使用 python-docx 提取文本：
- 段落：保留标题层级标记（# 前缀，与 Markdown 分块器兼容）
- 表格：按行展开为「列1 | 列2 | 列3」文本
"""

from pathlib import Path

from app.core.logger import get_logger
from app.parsers.base import BaseParser

logger = get_logger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# 标题样式名 → Markdown 级别
_HEADING_LEVELS = {
    "Heading 1": "#",
    "Heading 2": "##",
    "Heading 3": "###",
    "Heading 4": "####",
    "Title": "#",
}


class DocxParser(BaseParser):
    """Word 文档解析器（.docx，不支持旧版 .doc）"""

    def parse(self, file_path: str | Path) -> str:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            import docx

            document = docx.Document(str(path))
        except Exception as e:
            raise ValueError(f"Word 文件无法打开，可能已损坏或不是 .docx 格式: {e}") from e

        parts: list[str] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            level = _HEADING_LEVELS.get(para.style.name if para.style else "")
            parts.append(f"{level} {text}" if level else text)

        # 表格：按行展开，单元格以 | 分隔
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        result = "\n\n".join(parts)
        if not result.strip():
            logger.warning(f"Word 未提取到文字内容: {path.name}")

        logger.info(f"Word 解析完成: {path.name}, 提取 {len(result)} 字符")
        return result

    @property
    def supported_formats(self) -> list[str]:
        return [DOCX_MIME]
